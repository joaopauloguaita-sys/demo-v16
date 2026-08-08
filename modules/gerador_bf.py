import os
import sys
from datetime import date, datetime
from tkinter import messagebox

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Caminho do modelo oficial (o BF.docx da prefeitura, já com a tabela pronta)
# Sobe dois níveis porque este arquivo fica dentro de 'modules', e 'assets' fica na raiz do projeto
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(BASE_DIR, "assets", "BF_modelo.docx")
SAIDA_DIR = os.path.join(BASE_DIR, "relatorios_gerados")

# Mesma função de conexão usada pelo resto do sistema (garante que sempre
# aponta para o banco de dados correto, o mesmo que a ficha usa para salvar)
sys.path.insert(0, BASE_DIR)
from database.db import get_connection


def calcular_periodo(mes_referencia, ano_referencia):
    """Período do BF: dia 16 do mês anterior até dia 15 do mês/ano escolhidos."""
    mes = int(mes_referencia)
    ano = int(ano_referencia)
    data_fim = date(ano, mes, 15)
    if mes == 1:
        mes_ini, ano_ini = 12, ano - 1
    else:
        mes_ini, ano_ini = mes - 1, ano
    data_ini = date(ano_ini, mes_ini, 16)
    return data_ini, data_fim


def buscar_profissionais(cursor):
    """Professores + Funcionários ativos, exceto Terceirizado, em ordem alfabética."""
    query = """
        SELECT id, nome, matricula, situacao_funcional, 'professores' as tabela
        FROM professores WHERE ativo = 1
        UNION ALL
        SELECT id, nome, matricula, situacao_funcional, 'funcionarios' as tabela
        FROM funcionarios WHERE ativo = 1
    """
    linhas = cursor.execute(query).fetchall()
    filtrados = [
        r for r in linhas
        if (r["situacao_funcional"] or "").strip().lower() != "terceirizado"
    ]
    filtrados.sort(key=lambda r: (r["nome"] or "").upper())
    return filtrados


def buscar_atestados(cursor, tabela, pessoa_id, data_ini, data_fim):
    query = """
        SELECT tipo, data, duracao, unidade_duracao FROM atestados
        WHERE entidade = ? AND entidade_id = ? AND data >= ? AND data <= ?
              AND (excluido IS NULL OR excluido = 0)
        ORDER BY data
    """
    return cursor.execute(
        query, (tabela, pessoa_id, data_ini.isoformat(), data_fim.isoformat())
    ).fetchall()


def _preencher_celula(cell, texto):
    """Escreve o texto na célula reaproveitando o parágrafo/estilo já existente no modelo."""
    p = cell.paragraphs[0]
    texto = "" if texto is None else str(texto)
    if p.runs:
        p.runs[0].text = texto
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(texto)


def _substituir_paragrafo(paragrafo, texto_novo, centralizar=False):
    if paragrafo.runs:
        paragrafo.runs[0].text = texto_novo
        for run in paragrafo.runs[1:]:
            run.text = ""
    else:
        paragrafo.add_run(texto_novo)
    if centralizar:
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _formatar_data(data_str):
    try:
        return datetime.strptime(data_str, "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return data_str or "-"


def gerar_bf(mes_referencia, ano_referencia, nome_diretora, parent=None):
    if not os.path.exists(TEMPLATE_PATH):
        messagebox.showerror(
            "Erro",
            f"Modelo do BF não encontrado em:\n{TEMPLATE_PATH}\n\n"
            "Coloque o arquivo BF_modelo.docx dentro da pasta 'assets'.",
            parent=parent,
        )
        return None

    data_ini, data_fim = calcular_periodo(mes_referencia, ano_referencia)

    conn = get_connection()
    cursor = conn.cursor()
    profissionais = buscar_profissionais(cursor)

    # Município atual, direto do cadastro em Dados da Escola — evita depender
    # de um texto fixo (ex: "Cambará,") que pode não bater mais com o modelo
    escola_row = cursor.execute("SELECT municipio FROM dados_escola LIMIT 1").fetchone()
    municipio_atual = (escola_row["municipio"] if escola_row and escola_row["municipio"] else "").strip()

    doc = Document(TEMPLATE_PATH)
    tabela_principal = doc.tables[0]
    tabela_atestados = doc.tables[1]

    # Linha 0 = cabeçalho, linha 1 = legenda (X-X-X...) — dados começam na linha 2
    linhas_principais = tabela_principal.rows[2:]
    # Linha 0 = cabeçalho (Nome/Duração/Data) — dados começam na linha 1
    linhas_atestados = tabela_atestados.rows[1:]

    avisos = []
    idx_atestado = 0

    for i, pessoa in enumerate(profissionais):
        if i >= len(linhas_principais):
            avisos.append(
                f"O modelo só tem {len(linhas_principais)} linhas na tabela principal, "
                f"mas há {len(profissionais)} profissionais. Os que não couberam ficaram de fora."
            )
            break

        linha = linhas_principais[i]
        atestados = buscar_atestados(cursor, pessoa["tabela"], pessoa["id"], data_ini, data_fim)

        _preencher_celula(linha.cells[0], pessoa["matricula"] or "-")
        _preencher_celula(linha.cells[1], (pessoa["nome"] or "").upper())

        if atestados:
            _preencher_celula(linha.cells[7], "X")  # coluna FALTA -> Just.
            tipos = ", ".join(sorted({(a["tipo"] or "").strip() for a in atestados if a["tipo"]}))
            _preencher_celula(linha.cells[9], tipos)  # coluna de observação

            for a in atestados:
                if idx_atestado >= len(linhas_atestados):
                    avisos.append(
                        f"O modelo só tem {len(linhas_atestados)} linhas na tabela de "
                        "Atestado/Declaração — alguns atestados não couberam."
                    )
                    break
                linha_at = linhas_atestados[idx_atestado]
                duracao_txt = f"{a['duracao'] or ''} {a['unidade_duracao'] or ''}".strip()
                _preencher_celula(linha_at.cells[0], (pessoa["nome"] or "").upper())
                _preencher_celula(linha_at.cells[1], duracao_txt)
                _preencher_celula(linha_at.cells[2], _formatar_data(a["data"]))
                idx_atestado += 1

    conn.close()

    # Parágrafos fora das tabelas: período, data de emissão e assinatura
    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto.startswith("Período:"):
            _substituir_paragrafo(
                p, f"Período: {data_ini.strftime('%d/%m/%Y')} à {data_fim.strftime('%d/%m/%Y')}"
            )
        elif texto.startswith("Cambará,") or (municipio_atual and texto.startswith(f"{municipio_atual},")):
            cidade_txt = municipio_atual or "Cambará"
            _substituir_paragrafo(p, f"{cidade_txt}, {date.today().strftime('%d/%m/%Y')}")
        elif "ASS. DIRETOR" in texto.upper():
            _substituir_paragrafo(p, f"ASS. {nome_diretora.upper()}", centralizar=True)

    os.makedirs(SAIDA_DIR, exist_ok=True)
    nome_arquivo = f"BF_{int(mes_referencia):02d}_{ano_referencia}.docx"
    caminho_completo = os.path.join(SAIDA_DIR, nome_arquivo)

    try:
        doc.save(caminho_completo)
    except PermissionError:
        messagebox.showerror(
            "Arquivo em uso",
            f"Não consegui salvar porque o arquivo abaixo parece estar aberto "
            f"(provavelmente no Word, de uma geração anterior):\n\n{caminho_completo}\n\n"
            "Feche esse arquivo e clique em Gerar novamente.",
            parent=parent,
        )
        return None

    if avisos:
        messagebox.showwarning("Atenção", "\n\n".join(avisos), parent=parent)

    try:
        os.startfile(caminho_completo)
    except Exception as e:
        messagebox.showerror(
            "Gerado, mas não abriu sozinho",
            f"O arquivo foi salvo em:\n{caminho_completo}\n\n"
            f"Mas não consegui abrir automaticamente:\n{e}",
            parent=parent,
        )

    return caminho_completo


# Mantido por compatibilidade com quem já chama a função pelo nome antigo
def gerar_pdf_bf(mes_referencia, ano_referencia, nome_diretora, parent=None):
    return gerar_bf(mes_referencia, ano_referencia, nome_diretora, parent=parent)
